from __future__ import annotations

import sys
import tempfile
import unittest
import zipfile
from pathlib import Path

from docx import Document
from lxml import etree


ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "scripts"))

from build_report_from_template import build_report
from generate_submission_docs import make_user_guide


CORE_NS = {
    "dc": "http://purl.org/dc/elements/1.1/",
    "cp": "http://schemas.openxmlformats.org/package/2006/metadata/core-properties",
}


def office_authors(path: Path) -> tuple[str, str]:
    with zipfile.ZipFile(path) as archive:
        root = etree.fromstring(archive.read("docProps/core.xml"))
    creator = root.findtext("dc:creator", default="", namespaces=CORE_NS)
    modified_by = root.findtext("cp:lastModifiedBy", default="", namespaces=CORE_NS)
    return creator, modified_by


class SubmissionQualityTests(unittest.TestCase):
    def test_report_is_built_from_retained_template(self):
        values = {
            "{{DETECTION_RATE}}": "97.00%",
            "{{FPR}}": "15.33%",
            "{{PRECISION}}": "67.83%",
            "{{F1}}": "0.7984",
            "{{AUC}}": "0.9708",
            "{{LATENCY_MS}}": "0.144",
            "{{THROUGHPUT}}": "6947.3",
            "{{MODEL_MB}}": "0.303",
            "{{SEQUENCE_COUNT}}": "400",
            "{{PARAMETERS}}": "67,522",
        }
        with tempfile.TemporaryDirectory() as directory:
            output = Path(directory, "report.docx")
            build_report(output, values)
            doc = Document(output)
            text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
            footer_text = "\n".join(
                paragraph.text
                for section in doc.sections
                for paragraph in section.footer.paragraphs
            )
            with zipfile.ZipFile(output) as archive:
                footer_xml = "\n".join(
                    archive.read(name).decode("utf-8")
                    for name in archive.namelist()
                    if name.startswith("word/footer") and name.endswith(".xml")
                )
                settings_xml = archive.read("word/settings.xml").decode("utf-8")
            self.assertEqual(len(doc.sections), 9)
            self.assertIn("作品类型：开放式自由命题", text)
            self.assertIn("Transformer 单模型", text)
            self.assertIn("实际混合告警链路", text)
            self.assertNotIn("填写说明", text)
            self.assertNotIn("共 7页", footer_text)
            self.assertIn("PAGE", footer_xml)
            self.assertIn("NUMPAGES", footer_xml)
            self.assertIn("updateFields", settings_xml)
            for phrase in (
                "Agent 与主机双域联合建模",
                "模型分数与透明安全规则双证据融合",
                "困难负样本与场景留出",
            ):
                self.assertIn(phrase, text)
            self.assertEqual(office_authors(output), ("", ""))
            self.assertNotIn("（建议包括", text)
            self.assertNotIn("（请简要说明", text)
            self.assertGreaterEqual(len(doc.tables), 2)
            self.assertGreaterEqual(len(doc.inline_shapes), 3)

    def test_user_guide_contains_four_visual_examples_and_no_author(self):
        with tempfile.TemporaryDirectory() as directory:
            output = Path(directory, "guide.docx")
            make_user_guide(output)
            doc = Document(output)
            text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
            self.assertGreaterEqual(len(doc.inline_shapes), 4)
            for phrase in (
                "系统启动与主界面",
                "日志输入示例",
                "告警结果页面",
                "模型分、规则分与证据链",
            ):
                self.assertIn(phrase, text)
            self.assertEqual(office_authors(output), ("", ""))

    def test_final_office_deliverables_do_not_expose_authors(self):
        required = (
            ROOT / "submission" / "作品报告_AgentGuard_自主命题模板版.docx",
            ROOT / "submission" / "系统使用说明_AgentGuard.docx",
            ROOT / "submission" / "作品原创性声明_待签字盖章.docx",
        )
        optional = ROOT / "submission" / "答辩PPT_AgentGuard_自主命题最终版.pptx"
        for path in required:
            self.assertTrue(path.is_file(), f"提交物缺失：{path.name}")
            self.assertEqual(office_authors(path), ("", ""), path.name)
        if optional.is_file():
            self.assertEqual(office_authors(optional), ("", ""), optional.name)

    def test_authored_materials_do_not_claim_directed_track_or_organizer_data(self):
        paths = [ROOT / "README.md", *sorted((ROOT / "docs").glob("0[1-8]_*.md"))]
        banned = (
            "本项目对应官方定向式",
            "定向式专项命题（题目3）",
            "组委会正式数据指标",
            "官方数据到手后",
            "为官方 20% 上限",
        )
        for path in paths:
            text = path.read_text(encoding="utf-8")
            for phrase in banned:
                self.assertNotIn(phrase, text, f"{path.name} 仍含错误赛道口径")

    def test_generator_preserves_independently_built_ppt(self):
        source = (ROOT / "scripts" / "generate_submission_docs.py").read_text(
            encoding="utf-8"
        )
        main_source = source[source.index("def main():") :]
        self.assertIn("build_report(official_report, values)", main_source)
        self.assertNotIn("make_ppt(", main_source)


if __name__ == "__main__":
    unittest.main()
