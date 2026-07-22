from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from _bootstrap import ROOT
from office_hygiene import scrub_office_metadata


TEMPLATE = ROOT / "00_比赛原始材料" / "work-report-template.docx"
REPORT_SOURCE = ROOT / "docs" / "05_作品报告正文.md"
SECTION_TITLES = (
    "摘要",
    "第一章 作品概述",
    "第二章 作品设计与实现",
    "第三章 作品测试与分析",
    "第四章 创新性说明",
    "第五章 总结",
    "参考文献",
)


def _set_run_font(run, name: str = "宋体", size: float = 12, bold=None) -> None:
    run.font.name = "Times New Roman" if name == "宋体" else name
    run.font.size = Pt(size)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), name)
    if bold is not None:
        run.bold = bold


def _clean_inline(text: str) -> str:
    return re.sub(r"[`*_]", "", text).strip()


def _format_body(paragraph, *, indent: bool = True) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.first_line_indent = Cm(0.74) if indent else Cm(0)
    for run in paragraph.runs:
        _set_run_font(run)


def _add_before(doc: Document, anchor, text: str = "", style: str | None = None):
    paragraph = doc.add_paragraph(text, style=style)
    anchor.addprevious(paragraph._p)
    return paragraph


def _remove_between(heading, anchor) -> None:
    element = heading._p.getnext()
    while element is not None and element is not anchor:
        following = element.getnext()
        element.getparent().remove(element)
        element = following


def _remove_instruction_section(doc: Document) -> None:
    instruction = next(
        (paragraph for paragraph in doc.paragraphs if paragraph.text.strip() == "填写说明"),
        None,
    )
    if instruction is None:
        return
    element = instruction._p
    removed_section_break = False
    while element is not None:
        following = element.getnext()
        properties = getattr(element, "pPr", None)
        has_section_break = properties is not None and properties.sectPr is not None
        element.getparent().remove(element)
        if has_section_break:
            removed_section_break = True
            break
        element = following
    if not removed_section_break:
        raise ValueError("填写说明节缺少结束分节符，拒绝生成可能损坏的报告")


def _parse_sections(markdown: str) -> dict[str, list[str]]:
    sections: dict[str, list[str]] = {}
    current = None
    for raw_line in markdown.splitlines():
        if raw_line.startswith("# "):
            current = raw_line[2:].strip()
            sections[current] = []
        elif current is not None:
            sections[current].append(raw_line)
    missing = set(SECTION_TITLES) - set(sections)
    if missing:
        raise ValueError(f"报告源缺少章节：{sorted(missing)}")
    return sections


def _set_table_geometry(table, widths_cm: list[float]) -> None:
    table.autofit = False
    table.alignment = 1
    total_twips = int(sum(widths_cm) * 567)
    table_properties = table._tbl.tblPr
    width = table_properties.first_child_found_in("w:tblW")
    if width is None:
        width = OxmlElement("w:tblW")
        table_properties.append(width)
    width.set(qn("w:type"), "dxa")
    width.set(qn("w:w"), str(total_twips))

    indent = table_properties.first_child_found_in("w:tblInd")
    if indent is None:
        indent = OxmlElement("w:tblInd")
        table_properties.append(indent)
    indent.set(qn("w:type"), "dxa")
    indent.set(qn("w:w"), "120")

    borders = table_properties.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table_properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{edge}"))
        if border is None:
            border = OxmlElement(f"w:{edge}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), "808080")

    for row_index, row in enumerate(table.rows):
        if row_index == 0:
            row_properties = row._tr.get_or_add_trPr()
            repeat = OxmlElement("w:tblHeader")
            repeat.set(qn("w:val"), "true")
            row_properties.append(repeat)
        for column_index, cell in enumerate(row.cells):
            cell.width = Cm(widths_cm[column_index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell_properties = cell._tc.get_or_add_tcPr()
            margins = cell_properties.first_child_found_in("w:tcMar")
            if margins is None:
                margins = OxmlElement("w:tcMar")
                cell_properties.append(margins)
            for side, value in (("top", 100), ("left", 120), ("bottom", 100), ("right", 120)):
                node = margins.find(qn(f"w:{side}"))
                if node is None:
                    node = OxmlElement(f"w:{side}")
                    margins.append(node)
                node.set(qn("w:w"), str(value))
                node.set(qn("w:type"), "dxa")


def _add_table(doc: Document, anchor, rows: list[list[str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    anchor.addprevious(table._tbl)
    if len(rows[0]) == 2:
        widths = [9.0, 6.0]
    elif len(rows[0]) == 5:
        widths = [4.8, 2.55, 2.55, 2.55, 2.55]
    else:
        widths = [15.0 / len(rows[0])] * len(rows[0])
    _set_table_geometry(table, widths)
    for row_index, values in enumerate(rows):
        for column_index, value in enumerate(values):
            cell = table.cell(row_index, column_index)
            cell.text = _clean_inline(value)
            if row_index == 0:
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), "E7EEF8")
                cell._tc.get_or_add_tcPr().append(shading)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
                if row_index == 0 or column_index > 0
                else WD_ALIGN_PARAGRAPH.LEFT
            )
            paragraph.paragraph_format.line_spacing = 1.2
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                _set_run_font(run, size=10.5, bold=(row_index == 0))
    spacer = _add_before(doc, anchor)
    spacer.paragraph_format.space_after = Pt(3)


def _add_image(doc: Document, anchor, alt_text: str, relative_path: str, figure_number: int) -> None:
    path = ROOT / relative_path
    if not path.exists():
        raise FileNotFoundError(f"报告插图不存在：{path}")
    paragraph = _add_before(doc, anchor)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.add_run().add_picture(str(path), width=Cm(15.0))
    caption = _add_before(doc, anchor, f"图 {figure_number}  {alt_text}")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(6)
    for run in caption.runs:
        _set_run_font(run, size=10.5)


def _write_section(doc: Document, anchor, lines: list[str], figure_number: int) -> int:
    index = 0
    while index < len(lines):
        line = lines[index].strip()
        if not line:
            index += 1
            continue
        if line.startswith("|") and index + 1 < len(lines) and re.match(
            r"^\|?\s*:?-+", lines[index + 1].strip()
        ):
            rows = [[cell.strip() for cell in line.strip("|").split("|")]]
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                rows.append([
                    cell.strip() for cell in lines[index].strip().strip("|").split("|")
                ])
                index += 1
            _add_table(doc, anchor, rows)
            continue
        image = re.fullmatch(r"!\[(.+?)]\((.+?)\)", line)
        if image:
            figure_number += 1
            _add_image(doc, anchor, image.group(1), image.group(2), figure_number)
            index += 1
            continue
        if line.startswith("## "):
            paragraph = _add_before(doc, anchor, _clean_inline(line[3:]), "Heading 2")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_before = Pt(12)
            paragraph.paragraph_format.space_after = Pt(6)
            for run in paragraph.runs:
                _set_run_font(run, "黑体", 14, True)
        elif line.startswith("### "):
            paragraph = _add_before(doc, anchor, _clean_inline(line[4:]), "Heading 3")
            for run in paragraph.runs:
                _set_run_font(run, "黑体", 12, True)
        elif re.match(r"^[-*] ", line):
            paragraph = _add_before(doc, anchor, _clean_inline(line[2:]), "List Bullet")
            _format_body(paragraph, indent=False)
        elif re.match(r"^\d+\. ", line):
            paragraph = _add_before(
                doc, anchor, _clean_inline(re.sub(r"^\d+\. ", "", line)), "List Number"
            )
            _format_body(paragraph, indent=False)
        else:
            paragraph = _add_before(doc, anchor, _clean_inline(line))
            _format_body(paragraph)
        index += 1
    return figure_number


def _replace_cover(doc: Document) -> None:
    values = {
        "作品名称：": "作品名称：AgentGuard：基于行为序列建模的 AI Agent 异常检测系统",
        "作品类型：": "作品类型：开放式自由命题",
        "电子邮箱：": "电子邮箱：【提交前填写】",
        "提交日期：": "提交日期：2026 年【  】月【  】日",
    }
    for paragraph in doc.paragraphs:
        for prefix, value in values.items():
            if paragraph.text.startswith(prefix):
                paragraph.clear()
                run = paragraph.add_run(value)
                _set_run_font(run, "黑体", 16, True)
                paragraph.paragraph_format.line_spacing = 1.5
                break


def _enable_field_updates(doc: Document) -> None:
    settings = doc.settings._element
    updates = settings.findall(qn("w:updateFields"))
    if not updates:
        updates = [OxmlElement("w:updateFields")]
        settings.append(updates[0])
    for update in updates:
        update.set(qn("w:val"), "true")


def _append_page_field(paragraph, instruction: str) -> None:
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), instruction)
    field.set(qn("w:dirty"), "true")
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia"):
        fonts.set(qn(key), "宋体")
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "18")
    run_properties.extend((fonts, size))
    text = OxmlElement("w:t")
    text.text = "1"
    run.extend((run_properties, text))
    field.append(run)
    paragraph._p.append(field)


def _replace_page_footers(doc: Document) -> None:
    seen_parts = set()
    for section in doc.sections:
        footer = section.footer
        part_name = str(footer.part.partname)
        if part_name in seen_parts:
            continue
        seen_parts.add(part_name)
        for child in list(footer._element):
            footer._element.remove(child)
        paragraph = footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run("第 ")
        _set_run_font(run, size=9)
        _append_page_field(paragraph, "PAGE")
        run = paragraph.add_run(" 页 共 ")
        _set_run_font(run, size=9)
        _append_page_field(paragraph, "NUMPAGES")
        run = paragraph.add_run(" 页")
        _set_run_font(run, size=9)


def _ensure_heading_styles(doc: Document) -> None:
    for level, size in ((2, 14), (3, 12)):
        name = f"Heading {level}"
        if name in doc.styles:
            style = doc.styles[name]
        else:
            style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = doc.styles["Normal"]
        style.font.name = "黑体"
        style.font.size = Pt(size)
        style.font.bold = True
        style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "黑体")
        properties = style._element.get_or_add_pPr()
        outline = properties.find(qn("w:outlineLvl"))
        if outline is None:
            outline = OxmlElement("w:outlineLvl")
            properties.append(outline)
        outline.set(qn("w:val"), str(level - 1))


def build_report(output_path: str | Path, replacements: dict[str, str]) -> Path:
    template_hash = "90BFFB30093F40080EBFFAD1D8E72345274C87F9C6595AFE6F2189B3C5D32568"
    import hashlib

    actual_hash = hashlib.sha256(TEMPLATE.read_bytes()).hexdigest().upper()
    if actual_hash != template_hash:
        raise ValueError("作品报告模板已变化，需要重新提炼模板")

    markdown = REPORT_SOURCE.read_text(encoding="utf-8")
    for token, value in replacements.items():
        markdown = markdown.replace(token, value)
    if re.search(r"\{\{[A-Z0-9_]+\}\}", markdown):
        raise ValueError("作品报告仍包含未替换的指标占位符")
    sections = _parse_sections(markdown)

    doc = Document(TEMPLATE)
    _ensure_heading_styles(doc)
    _replace_cover(doc)
    _remove_instruction_section(doc)
    headings = {paragraph.text.strip(): paragraph for paragraph in doc.paragraphs}
    for title in SECTION_TITLES:
        if title not in headings and not (
            title == "参考文献" and "参考文献 " in headings
        ):
            raise ValueError(f"模板缺少章节槽位：{title}")
    if "参考文献" not in headings:
        headings["参考文献"] = headings["参考文献 "]
        headings["参考文献"].text = "参考文献"

    figure_number = 0
    for title in SECTION_TITLES:
        heading = headings[title]
        anchor = None
        element = heading._p.getnext()
        while element is not None:
            properties = getattr(element, "pPr", None)
            if properties is not None and properties.sectPr is not None:
                anchor = element
                break
            element = element.getnext()
        if anchor is None:
            anchor = doc._body._element.sectPr
        _remove_between(heading, anchor)
        figure_number = _write_section(doc, anchor, sections[title], figure_number)

    _replace_page_footers(doc)
    _enable_field_updates(doc)
    doc.core_properties.title = "AgentGuard：基于行为序列建模的 AI Agent 异常检测系统"
    doc.core_properties.subject = "自主命题赛道作品报告"
    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""
    doc.core_properties.comments = ""
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    scrub_office_metadata(output)
    return output


if __name__ == "__main__":
    import argparse
    import json

    parser = argparse.ArgumentParser(description="基于官方模板生成 AgentGuard 作品报告")
    parser.add_argument("--output", required=True)
    parser.add_argument("--replacements", help="可选：指标占位符 JSON 文件")
    args = parser.parse_args()
    if args.replacements:
        values = json.loads(Path(args.replacements).read_text(encoding="utf-8"))
    else:
        metrics = json.loads(
            (ROOT / "artifacts" / "evaluation" / "metrics.json").read_text(encoding="utf-8")
        )
        training = json.loads(
            (ROOT / "artifacts" / "training_result.json").read_text(encoding="utf-8")
        )
        values = {
            "{{DETECTION_RATE}}": f"{metrics['detection_rate'] * 100:.2f}%",
            "{{FPR}}": f"{metrics['false_positive_rate'] * 100:.2f}%",
            "{{PRECISION}}": f"{metrics['precision'] * 100:.2f}%",
            "{{F1}}": f"{metrics['f1']:.4f}",
            "{{AUC}}": f"{metrics['roc_auc']:.4f}",
            "{{LATENCY_MS}}": f"{metrics['mean_latency_ms']:.3f}",
            "{{THROUGHPUT}}": f"{metrics['throughput_sequences_per_second']:.1f}",
            "{{MODEL_MB}}": f"{metrics['checkpoint_size_mb']:.3f}",
            "{{SEQUENCE_COUNT}}": str(metrics["sequence_count"]),
            "{{PARAMETERS}}": f"{training['model_parameters']:,}",
        }
    print(build_report(args.output, values))
